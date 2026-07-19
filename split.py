# -*- coding: utf-8 -*-
import os
import sys
import subprocess
import time

# 필수 라이브러리 체크 및 자동 설치 함수
def ensure_dependencies():
    dependencies = {
        "docx": "python-docx",
        "pypdf": "pypdf",
        "requests": "requests"
    }
    for module, package in dependencies.items():
        try:
            __import__(module)
        except ImportError:
            print(f"[*] 필수 라이브러리 '{package}'가 설치되어 있지 않습니다. 자동 설치를 시작합니다...")
            try:
                # 윈도우 환경을 위해 subprocess.run 또는 check_call 활용
                subprocess.check_call([sys.executable, "-m", "pip", "install", package])
                print(f"[+] '{package}' 설치가 성공적으로 완료되었습니다.\n")
            except Exception as e:
                print(f"[-] '{package}' 설치에 실패했습니다: {e}")
                print("수동으로 설치한 후 스크립트를 다시 실행해 주세요: pip install python-docx pypdf requests")
                sys.exit(1)

ensure_dependencies()

import requests
import docx
import pypdf

# 1. 텍스트 추출 모듈 구현
def extract_text_from_txt(file_path):
    encodings = ['utf-8', 'cp949', 'latin-1']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"지원하는 인코딩으로 파일을 읽을 수 없습니다: {file_path}")

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 테이블 표 내부에 있는 내용도 추출해서 텍스트 스트림에 포함
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append(" | ".join(row_text))
    return "\n".join(full_text)

def extract_text_from_pdf(file_path):
    reader = pypdf.PdfReader(file_path)
    full_text = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            full_text.append(text)
    return "\n".join(full_text)

def extract_text(file_path):
    _, ext = os.path.splitext(file_path.lower())
    if ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {ext}")

# 2. 문맥 단락 분절(Chunking) 모듈
def split_into_chunks(text, max_chunk_size=3500):
    paragraphs = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0

    for para in paragraphs:
        para_len = len(para) + 1  # 개행 문자 길이를 포함
        if current_length + para_len > max_chunk_size:
            if current_chunk:
                chunks.append("\n".join(current_chunk))
                current_chunk = [para]
                current_length = para_len
            else:
                # 개별 문단 하나의 크기가 max_chunk_size를 넘는 경우 예외적으로 단독 청크 처리
                chunks.append(para)
                current_chunk = []
                current_length = 0
        else:
            current_chunk.append(para)
            current_length += para_len

    if current_chunk:
        chunks.append("\n".join(current_chunk))

    return chunks

# 3. Ollama Gemma4 API 연동 설정 및 가공 함수
OLLAMA_URL = "http://127.0.0.1:11434/api/generate"
MODEL_NAME = "hf.co/lmstudio-community/gemma-4-E4B-it-GGUF:Q4_K_M"

# 1-36.txt 분석을 토대로 구조화한 정밀 LLM 지침서
SYSTEM_PROMPT = """You are an expert text formatter and editor. Your task is to clean up, format, and apply Markdown syntax to the provided raw book text according to the following strict guidelines:

1. Metadata Filtration (Front Matter):
   - Remove publisher ads, series lists, detailed copyright pages, and cataloging information. 
   - Retain only the title, author name, contents, and the main chapters/sections.

2. Markdown Heading Hierarchy:
   - Convert the book title into a level-1 heading (# Title).
   - Convert major sections (Contents, Preface, Acknowledgements, Introduction, Chapter Titles) into level-2 headings (## Section).
   - Convert sub-sections (e.g., FEMINISM) into level-3 headings (### Subsection).

3. Blockquotes formatting:
   - Identify extended quotations or excerpts from other literary works or studies.
   - Prepend "> " to every line of these quote blocks to format them as Markdown blockquotes.

4. Paragraph Reconstruction & Line Merging:
   - Merge sentences that were broken across multiple lines due to PDF/DOCX layout extraction errors. Join them into smooth, continuous paragraphs.
   - Do NOT modify or rewrite the actual words or spellings. Do not translate. Keep the original English text exactly as it is.
   - Ensure paragraphs are separated by exactly two empty lines (i.e. '\\n\\n\\n').

5. Page Number & Page Header Removal:
   - Identify and strip out orphan page numbers (e.g., "54", "iii") and running headers/footers that interrupt the text.
   - Do NOT remove numeric citation anchors embedded in the text (e.g., "man.5" or "(Atwood 1979:7)").

6. Output Format:
   - Output ONLY the formatted Markdown text. Do not include any introduction, notes, explanation, or conversational filler (like "Here is the formatted text:")."""

def process_chunk_with_llm(chunk, chunk_index, total_chunks):
    prompt = f"<start_of_turn>system\n{SYSTEM_PROMPT}<end_of_turn>\n<start_of_turn>user\n[Processing Chunk {chunk_index + 1}/{total_chunks}]\nHere is the raw text to format:\n\n{chunk}<end_of_turn>\n<start_of_turn>model\n"
    
    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": False,
        "options": {
            "temperature": 0.1,  # 텍스트 복원의 정합성을 위해 온도를 낮춤
            "top_p": 0.9
        }
    }
    
    try:
        response = requests.post(OLLAMA_URL, json=payload, timeout=180)
        response.raise_for_status()
        result_json = response.json()
        return result_json.get("response", "").strip()
    except Exception as e:
        print(f"\n[-] 청크 [{chunk_index + 1}/{total_chunks}] 처리 중 오류 발생: {e}")
        # 오류 발생 시 텍스트 유실 방지를 위해 원본 청크를 그대로 보존 반환
        return chunk

# 4. 사용완료된 GPU VRAM 즉각 언로드(방출) 함수
def unload_vram():
    print("[*] 작업이 완료되어 VRAM 모델 언로드(메모리 해제)를 요청합니다...")
    payload = {
        "model": MODEL_NAME,
        "prompt": "",
        "stream": False,
        "keep_alive": 0
    }
    try:
        requests.post(OLLAMA_URL, json=payload, timeout=10)
        print("[+] VRAM 메모리가 완전히 방출되었습니다.")
    except Exception as e:
        print(f"[-] VRAM 방출 요청 실패 (Ollama 가 바쁘거나 이미 닫혀 있을 수 있음): {e}")

# 5. CLI 실행 진입 메인 함수
def main():
    # 명령행 인자가 전달되지 않은 경우, 자동 파일 감지 대화식 인터페이스 실행
    if len(sys.argv) < 2:
        print("사용법: python split.py <가공할_파일명>")
        print("예시: python split.py \"Feminist Stylistics (Sara Mills).docx\"")
        
        supported_extensions = ('.docx', '.pdf', '.txt')
        files = [f for f in os.listdir('.') if f.endswith(supported_extensions) and not f.startswith('splitted_') and f != '1-36.txt']
        if not files:
            print("\n[-] 폴더 내에 처리할 수 있는 문서 파일(.docx, .pdf, .txt)이 존재하지 않습니다.")
            sys.exit(1)
        
        print("\n[+] 현재 작업 디렉토리에서 발견된 대상 파일 목록:")
        for idx, f in enumerate(files):
            print(f"  [{idx + 1}] {f}")
        
        try:
            choice = int(input("\n작업할 파일의 번호를 선택하세요 (종료하려면 Enter): ")) - 1
            if 0 <= choice < len(files):
                file_path = files[choice]
            else:
                print("잘못된 선택입니다. 종료합니다.")
                sys.exit(1)
        except (ValueError, IndexError):
            print("작업이 취소되었습니다. 종료합니다.")
            sys.exit(1)
    else:
        file_path = sys.argv[1]

    if not os.path.exists(file_path):
        print(f"[-] 지정한 파일을 찾을 수 없습니다: {file_path}")
        sys.exit(1)

    print(f"\n[*] 처리 대상 파일: {file_path}")
    print("[*] 문서 데이터 추출 중...")
    try:
        raw_text = extract_text(file_path)
        print(f"[+] 문서 텍스트 추출 완료! (총 {len(raw_text)}글자)")
    except Exception as e:
        print(f"[-] 문서 추출 실패: {e}")
        sys.exit(1)

    print("[*] 텍스트를 문맥 단락별(3,500자 한도)로 분절하는 중...")
    chunks = split_into_chunks(raw_text)
    total_chunks = len(chunks)
    print(f"[+] 분절 완료! (총 {total_chunks}개의 작업 단위 청크 생성)")

    print("\n[*] 로컬 LLM (Ollama Gemma4)을 호출하여 가공 작업을 순차 실행합니다.")
    print("    컴퓨터 사양 및 파일 크기에 따라 다소 시간이 소요될 수 있습니다.\n")

    processed_results = []
    
    for idx, chunk in enumerate(chunks):
        print(f"⏳ 청크 가공 중 [{idx + 1}/{total_chunks}]...", end="", flush=True)
        start_time = time.time()
        
        processed_text = process_chunk_with_llm(chunk, idx, total_chunks)
        processed_results.append(processed_text)
        
        elapsed = time.time() - start_time
        print(f" 완료! ({elapsed:.1f}초 소요)")

    # 3개의 개행 문자로 조각들을 논리적으로 병합
    final_output = "\n\n\n".join(processed_results)

    # 출력 파일명 규칙: splitted_원래의파일명.txt
    dir_name, base_name = os.path.split(file_path)
    name_without_ext, _ = os.path.splitext(base_name)
    output_filename = f"splitted_{name_without_ext}.txt"
    output_path = os.path.join(dir_name, output_filename) if dir_name else output_filename

    print(f"\n[*] 정제 및 분절 완료 텍스트 저장 중: {output_path}")
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_output)
        print(f"[+] 최종 결과물 저장 완료! (최종 텍스트 크기: {len(final_output)}자)")
    except Exception as e:
        print(f"[-] 저장 실패: {e}")

    # 리소스 누수 방지를 위한 언로드 호출
    unload_vram()
    print("\n[+] 모든 처리가 끝났습니다. 에디터에서 불러와서 번역 및 요약을 시작할 수 있습니다!")

if __name__ == '__main__':
    main()
